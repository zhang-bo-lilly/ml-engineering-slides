"""Build coe_onepager.docx with exact 11pt font and 0.6in margins."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height   = Inches(11)
section.page_width    = Inches(8.5)
section.top_margin    = Inches(1.25)
section.bottom_margin = Inches(1.25)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

FONT = "Aptos"
PT   = 11

def set_spacing(para, before=0, after=3, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_run(para, text, bold=False, size=PT, color=None):
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold      = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def section_head(text):
    p = doc.add_paragraph()
    set_spacing(p, before=9, after=2)
    r = add_run(p, text, bold=True, size=PT)
    r.font.color.rgb = RGBColor(16, 100, 60)
    return p

def body(text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=4, line=PT * 1.2)
    add_run(p, text)
    return p

def bullet(label, text):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, before=0, after=3, line=PT * 1.2)
    p.paragraph_format.left_indent       = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_run(p, label + " ", bold=True)
    add_run(p, text)
    return p

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(title, before=0, after=2)
add_run(title, "Scientific AI Compute Center of Excellence", bold=True, size=13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(sub, before=0, after=8)
add_run(sub, "A Joint Initiative for AIR and Digital Core", size=PT, color=(80, 80, 80))

# ── Section 1 ─────────────────────────────────────────────────────────────────
section_head("The Inflection Point")
body(
    "The language of scientific AI has shifted — agentic workflows, autonomous discovery, AI-driven pipelines. "
    "Beneath the vocabulary, the compute challenge is the same: executing a directed graph of heterogeneous models "
    "where physics-based simulations, ML inference, and generative AI steps run in sequence and in parallel, each "
    "with distinct resource requirements. What has changed is the scale of that graph and the caliber of hardware "
    "each node demands. Tech@Lilly has a near-term window to define how this compute is owned, governed, and "
    "delivered — before fragmented build-out becomes structurally entrenched."
)

# ── Section 2 ─────────────────────────────────────────────────────────────────
section_head("The Current Pattern")
body(
    "Several LRL business-area teams have independently stood up AI compute environments. Examples include "
    "Applied Intelligence for Discovery (AI4D) under Discovery Oncology and Data Foundry within Lilly Small "
    "Molecule Discovery — each investing at the AVP level. Data Foundry has gone further, building a full "
    "platform organization structured across four AVP-led pillars. "
    "This reflects genuine urgency, but creates compounding risks:"
)
bullet("Scalability ceiling:", "Platforms built on lower-tier cloud GPUs face a cost cliff as workloads move to H100, H200, and B300 hardware — expensive, scarce, and materially higher in unit cost. Business leaders anchored on early cloud experience do not yet see the cliff ahead.")
bullet("Operational fragility:", "At least one team operates at significant scale without HPC expertise. Patterns harmless at lower tiers — e.g., using GPU instances as data staging vehicles — become punishing at H100/B300 scale. Pivoting requires months of planning and retraining, not a budget decision.")
bullet("Duplicated investment:", "Each team building independently means duplicated tooling, duplicated staffing, and no shared learning.")

# ── Section 3 ─────────────────────────────────────────────────────────────────
section_head("The Cost Reality")
body(
    "On-premise compute economics are compelling at Lilly's scale. MagTrain was deployed for $6.8M capex; "
    "AWS Savings Plans equivalent reaches $14.8M over three years (2.2×) and $24.7M over five years (3.6×). "
    "LillyPod's five-year TCO is $150M. Matching its 1,016 B300 GPUs on AWS Savings Plans costs $247M over "
    "three years — already 1.6× the full five-year on-premise TCO — and $412M over five years (2.7×)."
)

# ── Section 4 ─────────────────────────────────────────────────────────────────
section_head("The Proposal")
body(
    "The answer is a Scientific AI Compute Center of Excellence, jointly led by AIR and Digital Core. "
    "Scientific AI compute is, at its core, an infrastructure problem: capacity planning, power, advanced "
    "hardware procurement, and the operational expertise to run it at scale. That is the mandate of "
    "Infrastructure & Operations, and this CoE is its natural expression. Business-area teams that have "
    "moved quickly and independently would find in the CoE not a constraint, but a foundation to build on. "
    "We believe this is the right moment to move, and we\u2019d welcome the chance to define what that "
    "looks like together."
)

doc.save("coe_onepager.docx")
print("saved coe_onepager.docx")
