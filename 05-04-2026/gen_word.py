import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

LILLY_RED = RGBColor(0xCC, 0x00, 0x00)
DARK_GRAY = RGBColor(0x1F, 0x1F, 0x1F)
MID_GRAY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GRAY_BG = RGBColor(0xF2, 0xF2, 0xF2)
TABLE_HEADER_BG = RGBColor(0x1F, 0x1F, 0x1F)
TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_STRIPE = RGBColor(0xF7, 0xF7, 0xF7)

FONT_NAME = "Arial"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def no_border(cell):
    for side in ["top", "bottom", "left", "right"]:
        set_cell_borders(cell, **{side: {"val": "nil"}})


def add_run(para, text, bold=False, italic=False, size=10, color=None, font=FONT_NAME):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.keep_with_next = True
    if level == 1:
        run = para.add_run(text.upper())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(13)
        run.font.color.rgb = LILLY_RED
        # Red bottom border via paragraph border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "CC0000")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run = para.add_run(text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
    elif level == 3:
        run = para.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = MID_GRAY
    return para


def add_body(doc, text, italic=False, size=10, space_before=0, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    # Handle inline bold (**text**)
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            add_run(para, part[2:-2], bold=True, italic=italic, size=size, color=DARK_GRAY)
        else:
            add_run(para, part, italic=italic, size=size, color=MID_GRAY)
    return para


def add_bullet(doc, text, level=0, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    import re
    if bold_prefix:
        add_run(para, bold_prefix + " ", bold=True, size=10, color=DARK_GRAY)
        add_run(para, text, size=10, color=MID_GRAY)
    else:
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                add_run(para, part[2:-2], bold=True, size=10, color=DARK_GRAY)
            else:
                add_run(para, part, size=10, color=MID_GRAY)
    return para


def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = TABLE_HEADER_FG

    # Data rows
    import re
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = TABLE_STRIPE if r_idx % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            is_bold_row = cell_text.startswith("**") and cell_text.endswith("**")
            if is_bold_row:
                cell_text = cell_text[2:-2]
                run = p.add_run(cell_text)
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY
            else:
                parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK_GRAY

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_note(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            add_run(para, part[2:-2], bold=True, italic=True, size=9, color=MID_GRAY)
        else:
            add_run(para, part, italic=True, size=9, color=MID_GRAY)
    return para


def build():
    doc = Document()

    # Tight margins — body width = 6.77"
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(9)

    # ── TITLE ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(10)
    title_run = title_para.add_run("Compute Consolidation Thesis")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = LILLY_RED

    # ── OPENING + RISKS ───────────────────────────────────────────────────────
    add_body(doc,
        "LillyPod evolves into an enterprise platform for scientific-AI workflows. MagTrain's "
        "GPUs, new simulation-class GPUs, and fat CPU nodes consolidate onto a single Weka fast "
        "storage fabric under Run:ai. md3 joins the shared namespace without forced migration. "
        "Future compute requests route to the platform rather than spawning new purpose-built "
        "environments. Two compounding risks make this time-sensitive rather than aspirational.",
        size=9, space_after=5)

    add_body(doc,
        "**Risk 1 — False cost narrative.** Compute demand within LRL is being met through island "
        "environments on commodity GPUs. In the early phase these appear cost-effective. But "
        "commodity-GPU islands have a ceiling, and on the other side is a severe cost problem: "
        "training-class and simulation-class GPUs on cloud are not cheap at scale, and capacity "
        "at that tier requires months of lead time and upfront commitment. The teams generating "
        "this demand are not surfacing the future cost trajectory. By the time each island hits "
        "its ceiling, the narrative will have already set — the cheap-GPU path worked, and on-prem "
        "was the expensive option. Countering that after the fact is significantly harder than "
        "making the case now.",
        size=9, space_before=2, space_after=4)

    add_body(doc,
        "**Risk 2 — Structural intake.** Without a deliberate change in how compute requests are "
        "handled, the island pattern continues regardless of what the platform builds. Research IT, "
        "as the team that has historically fielded LRL compute demand, is the natural on-ramp — "
        "and it now sits within AIR, under your authority. The change required is operational: new "
        "requests route to the platform rather than spawning independent environments. That makes "
        "this consolidation durable, not a one-time hardware deployment.",
        size=9, space_before=2, space_after=8)

    # ── WHAT WE CONSOLIDATE ───────────────────────────────────────────────────
    add_heading(doc, "What We Consolidate")

    add_table(doc,
        headers=["Hardware", "Action"],
        rows=[
            ["MagTrain GPUs",
             "Weka fabric integration"],
            ["256× RTX 6000 Pro Blackwell",
             "New purchase, LillyPod fabric member; physics simulation, 96GB VRAM"],
            ["16× Dell R7625/R7725 CPU fat nodes",
             "New purchase, LillyPod fabric member; 4K CPU cores, 32TB RAM"],
            ["Weka hot tier +2PB",
             "Expand 4PB → 6PB as additional nodes, not drive expansion"],
        ],
        col_widths=[2.3, 4.47])

    add_body(doc,
        "Every item above joins the Weka fabric as a member, not an NFS mount. NFS accesses "
        "the hot tier but does not deliver fabric-level throughput or enable data-local compute.",
        size=9, space_after=8)

    # ── WHAT WE KEEP DISTRIBUTED ──────────────────────────────────────────────
    add_heading(doc, "What We Keep Distributed")

    add_body(doc,
        "md3 stays on Grid Engine for now. The forward step is Weka fabric membership — "
        "non-disruptive, no workload migration required.",
        size=9, space_after=8)

    # ── COST AND CAPABILITY CASE ──────────────────────────────────────────────
    add_heading(doc, "Cost and Capability Case")

    add_table(doc,
        headers=["Bucket", "Capital (est.)"],
        rows=[
            ["md3 — 256 RTX 6000 Pro (32 servers; expands md3 to 512)", "~$5.7M"],
            ["LillyPod — 256 RTX 6000 Pro + 16 CPU fat nodes + network/cables + Weka +2PB", "~$10.65M"],
            ["Shared DC infrastructure (UPS, racks, PDU, CDU, fire suppression)", "~$7.1M"],
            ["**2026 total**", "**~$23.45M**"],
        ],
        col_widths=[5.27, 1.5])

    add_body(doc, "**On-prem vs. cloud — LillyPod ask ($10.65M):**",
        size=9, space_before=4, space_after=2)

    add_table(doc,
        headers=["Line item", "On-prem", "Cloud equivalent", "Advantage"],
        rows=[
            ["256× RTX 6000 Pro", "~$5.7M", "~$19.2M (5-yr, AWS SP + 22% discount)", "**3.4×**"],
            ["Weka +2PB", "~$3M", "~$10M (3-yr FSx for Lustre at $0.14/GB/mo)", "**3.3×**"],
            ["CPU fat nodes", "~$1.3M", "Comparable or higher (high-memory EC2 families)", "—"],
            ["Network + fabric cabling", "~$650K", "No cloud equivalent", "—"],
        ],
        col_widths=[1.6, 0.7, 3.5, 0.97])

    add_body(doc,
        "To get Savings Plan rates, Lilly commits the full 3-year cost upfront on day one — "
        "access, not ownership. On-prem buys ownership and a shared fabric-connected platform "
        "that cloud cannot replicate.",
        size=9, space_after=4)

    add_body(doc,
        "**Live example.** The VS team runs simulation on cloud at $5.25/hr per job (L40S) and "
        "$3.63/hr (RTX 6000 Pro), at on-demand rates — the actual gap exceeds 7×. Jobs are "
        "hitting memory limits and splitting across GPUs, multiplying cost. H200 (141GB) and "
        "B300 (192GB) on the consolidated platform eliminate the split.",
        size=9, space_after=8)

    # ── MIGRATION PATH ────────────────────────────────────────────────────────
    add_heading(doc, "Migration Path")

    add_table(doc,
        headers=["Step", "Timing", "Outcome"],
        rows=[
            ["Step 0 — md3 RTX 6000 Pro deployment", "Near-term, in progress", "256× GPUs to md3; Grid Engine"],
            ["Step 1 — MagTrain Weka fabric integration", "Q3 2026", "H100/H200/L40S join LillyPod under Run:ai"],
            ["Step 2 — CPU fat node deployment", "Q3 2026", "Bioinformatics workloads on platform"],
            ["Step 3 — RTX 6000 Pro deployment", "Q4 2026", "End-to-end scientific-AI workflows on one fabric"],
            ["Step 4 — Weka +2PB expansion", "Q4 2026", "Storage proportional to GPU capacity increase"],
            ["Step 5 — md3 Weka fabric membership", "Q1 2027", "Data silo eliminated; md3 remains on Grid Engine"],
        ],
        col_widths=[2.4, 1.3, 3.07])

    # ── RISKS ─────────────────────────────────────────────────────────────────
    add_heading(doc, "Risks")

    add_table(doc,
        headers=["Risk", "Severity", "Mitigation"],
        rows=[
            ["MagTrain inter-row cabling physically constrained",
             "High", "Validate with Greg Johnson + Jonathan Klinginsmith before committing Step 1"],
            ["Weka expansion as drive-only creates throughput bottleneck",
             "High", "Specify additional nodes in procurement"],
            ["md3 fabric membership stalls at NFS state",
             "High", "Fabric membership must be explicit in the forward plan"],
            ["LRL-facing team continues fielding requests independently",
             "High", "Requires leadership direction; platform-first intake is not the default without it"],
        ],
        col_widths=[2.4, 0.6, 3.77])

    # ── WHAT I NEED FROM YOU ──────────────────────────────────────────────────
    add_heading(doc, "What I Need From You")

    asks = [
        ("1. Endorse end-to-end workflow throughput as the unit of optimization. ",
         "The hardware mix is only coherent if the goal is complete workflows on the platform. "
         "Confirm before the capital request goes forward."),
        ("2. Sponsor the ~$23.45M capital request ",
         "and carry it through the funding process."),
        ("3. Commit to cross-org coordination with Krishna Jagannathan. ",
         "Greg Johnson and Jonathan Klinginsmith own execution; a shared commitment at your level "
         "removes the coordination tax from the working level."),
        ("4. Direct the md3 storage fabric step. ",
         "Fabric membership — non-disruptive, no workload migration required."),
        ("5. Establish platform-first intake as a standing principle. ",
         "New LRL compute requests route through the platform. The Research IT merger creates "
         "the opportunity; this direction makes it durable."),
    ]

    for bold_text, body_text in asks:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        add_run(para, bold_text, bold=True, size=9, color=DARK_GRAY)
        add_run(para, body_text, size=9, color=MID_GRAY)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    out_path = os.path.join(BASE_DIR, "Compute Consolidation Plan.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
